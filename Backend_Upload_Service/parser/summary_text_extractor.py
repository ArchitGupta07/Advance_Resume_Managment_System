# Summary Extraction Flow
# Summary If Extracted good
# If summary has less than 25 words
# Name is extracted, and I search near name
# I may use words like result-driven and so on to match the para or I am, experience of, experience with
# Summary may not be there

# **********************
    #IF WITHIN 20 Words from back ". " (dot+space) is there we can remove those words 
 #*******************


# from moon_dummy import toExtract_Experience_Rolewise


from numpy import convolve
from pdfminer.converter import TextConverter
from pdfminer.pdfinterp import PDFPageInterpreter
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from docx import Document
# from copy_main_func import process_get_name

import spacy
from spacy.matcher import Matcher
import io
import requests
import re

def extract_text_from_pdf(pdf_path):

    
    if pdf_path.startswith('http'):
        response = requests.get(pdf_path)
        pdf_data = response.content
    else:
        with open(pdf_path, 'rb') as fh:
            pdf_data = fh.read()

    with io.BytesIO(pdf_data) as fh:
        for page in PDFPage.get_pages(fh, caching=True, check_extractable=True):
            resource_manager = PDFResourceManager()
            fake_file_handle = io.StringIO()
            converter = TextConverter(resource_manager, fake_file_handle, laparams=LAParams())
            page_interpreter = PDFPageInterpreter(resource_manager, converter)
            page_interpreter.process_page(page)
            text = fake_file_handle.getvalue()
            yield text
            converter.close()
            fake_file_handle.close()

def extract_from_docx(docx_file_path):
    doc = Document(docx_file_path)
    text = []
    for para in doc.paragraphs:
        text.append(para.text)

    tables_data = []
    for table in doc.tables:
        table_data = []  # List to hold all rows in a table
        for row in table.rows:
            row_data = []  # List to hold cell values in a row
            for cell in row.cells:
                row_data.append(cell.text)
            table_data.append(row_data)  # Append row data to table data
        tables_data.append(table_data)  # Append table data to tables data

        # Add table data to text
        table_text = '\n'.join(['\t'.join(row) for row in table_data])
        text.append(table_text)

    # return '\n'.join(text), tables_data
    return "\n".join(text)


from tika import parser
def tika_text_extraction(file_path):   
    parsed_file = parser.from_file(file_path)
    text = parsed_file["content"]
    return text

from collections import OrderedDict
def divide_resume_sections(resume_text):
    sections = OrderedDict()
    first = ""
    first_flag = True
  
    section_patterns = {

        "Summary":r"ABOUT|ABOUT ME|About Me|Overview|OVERVIEW|Summary|SUMMARY|summary|overview|Objective|OBJECTIVE|PROFILE|Profile",
        "Education": r"Education|EDUCATION",
        "Experience": r"\b(Experience|EXPERIENCE|Training|TRAINING|TRAININGS|INTERNSHIPS)\b",
        # "Training":r"Training|TRAINING|TRAININGS|INTERNSHIPS|INTERNSHIPS"        
        "Projects": r"Projects|PROJECTS",
        "Skills":r"SKILLS",
        "Achievements":"Achievements|ACHIEVEMENTS"
     
    }
    for sec in section_patterns:
        sections[sec]=""


    sect = {"Summary":["ABOUT","Overview","OVERVIEW","Summary","SUMMARY","summary","overview","Objective","OBJECTIVE","PROFILE","Profile"],"Education":["Education","EDUCATION"],"Experience":["Experience","EXPERIENCE","Training","TRAINING","TRAININGS","INTERNSHIPS"],"Projects":["Projects","PROJECTS"],"Skills":["SKILLS",""],"Achievements":["Achievements","ACHIEVEMENTS"]}
    
    for section_name in section_patterns:       
        matches = re.search(section_patterns[section_name], resume_text)     
        if matches:

            if first_flag:
                first = text[:matches.start()]
                first_flag=False
            # if section_name=="Summary":
            #     section_start = matches.start()    
                
            #     next_section_start = re.search(r"\b(ABOUT|Overview|OVERVIEW|Summary|SUMMARY|Education|EDUCATION|Experience|EXPERIENCE|Training|TRAINING|TRAININGS|INTERNSHIPS|Projects|PROJECTS|SKILLS|Achievements|ACHIEVEMENTS|$)\b(?!.*\d\+\syears\sof\sexperience)", resume_text[section_start+1:])          
            #     count=0
            #     check_end2 = 0
            #     final = 0
            # print(matches.group(0))
            section_start = matches.start()      
            next_section_start = re.search(r"\b(ABOUT|Overview|OVERVIEW|Summary|SUMMARY|Objective|OBJECTIVE|PROFILE|Profile|Education|EDUCATION|Experience|EXPERIENCE|Training|TRAINING|TRAININGS|INTERNSHIPS|Projects|PROJECT|SKILLS|Achievements|ACHIEVEMENTS|$)\b", resume_text[section_start+1:])          
            count=0
            check_end2 = 0
            final = 0

            # print(sec[section_name])
            if next_section_start:
                # print("xhexk ", next_section_start, next_section_start.group(0))
                while sect[section_name][0]==next_section_start.group(0) or sect[section_name][1]==next_section_start.group(0) :
                    check_end2 += section_start + next_section_start.end() +1 
                    next_section_start = re.search(r"\b(ABOUT|Overview|OVERVIEW|Summary|SUMMARY|Objective|OBJECTIVE|PROFILE|Profile|Education|EDUCATION|Experience|EXPERIENCE|Training|TRAINING|TRAININGS|INTERNSHIPS|Projects|PROJECTS|SKILLS|Achievements|ACHIEVEMENTS|$)\b", resume_text[check_end2:])    
                    if next_section_start:           
                        if sect[section_name][0]!=next_section_start.group(0) or sect[section_name][1]!=next_section_start.group(0):
                            final = next_section_start.start()+check_end2    
                    else:
                        break            
                    count+=1
                    if count==5:
                        break
                    flag=1  
            
            if next_section_start and next_section_start.group(0).replace(" ",""):                          
                section_end = section_start + next_section_start.start() +1              
            else:                
                section_end = len(resume_text)        
            if section_name in sections:                
                sections[section_name] += resume_text[section_start:section_end].strip()
            else:                
                sections[section_name] = resume_text[section_start:section_end].strip()

    return sections, first





# def extract_first_paragraph(text):
#     # Define a pattern to match the first 100 words of the first paragraph
#     pattern = re.compile(r'{1,100})')
    
#     # Search for the pattern in the text
#     match = pattern.search(text)
    
#     if match:
#         # Return the matched paragraph
#         return match.group().strip()
#     else:
#         return None

def extract_first_60_words(text):
    # Split the text into words
    words = text.strip().split()
    
    # Extract the first 60 words
    first_60_words = ' '.join(words[:60])
    
    return first_60_words


def extract_summary(resume_text, person_name):
    # Find the index of the person's name in the resume text
    name_index = resume_text.find(person_name)
    
    # Check if the name was found
    if name_index != -1:
        # Find the start of the summary section (right after the name)
        summary_start_index = name_index + len(person_name)
        
        # Find the end of the summary section (approximately 60-70 words later)
        summary_end_index = summary_start_index + 1  # Initialize with start of the summary
        word_count = 0
        while word_count < 70 and summary_end_index < len(resume_text):
            # Increment word count if a space is found
            if resume_text[summary_end_index] == ' ':
                word_count += 1
            summary_end_index += 1
        
        # Extract the summary section
        summary = resume_text[summary_start_index:summary_end_index]
        
        return summary.strip()  # Remove leading and trailing whitespace
        
    else:
        return "Person's name not found in the resume"
# text=""
# for page in extract_text_from_pdf("data/resumes/resume.pdf"):
#     text += ' ' + page 
# print(text)
# text = tika_text_extraction("data/resumes/resume.pdf")

# text ="""A data-driven Product Manager who has a proven track record of boosting revenue with 9 years of experience.                                        
# Experienced in tools like LLM, Midjourney, Stable diffusion, JIRA, Mixpanel, Google Analytics, A/B or multi 
# variate testing. On the business front corporate finance and Pricing strategy are my forte. Led the cross-
# functional teams in B2B, SaaS, B2C, marketing, sales, and service. Developed and demonstrated go-to-market 
# strategies, digital strategies, and product roadmaps. Successful KPIs include raising engagement levels, 
# # enhancing client satisfaction, and generating revenue through efficient tactics."""
# # print(text)

# text = tika_text_extraction("data/resumes/UDo.pdf")

# print(sections)

# print(text)
# filename = ''
# for i in range(2,20):
#     filename = f"resume{i}.pdf"
# a dummy strt




nlp = spacy.load("en_core_web_sm")

# matcher = Matcher(nlp.vocab)

def extract_prev_job_roles(text):
    job_roles = [
        "intern",
        "internship" "analyst",
        "developer",
        "manager",
        "engineer",
        "consultant",
        "designer",
        "specialist",
        "coordinator",
        "administrator",
        "executive",
        "assistant",
        "supervisor",
        "technician",
        "associate",
        "officer",
        "leader",
        "expert",
        "advisor",
        "strategist",
        "resources",
        "test",
    ]

    text = text.replace("\n", " ")
    # print(text)
    matcher = Matcher(nlp.vocab)
    pattern = [{"POS": "NOUN", "OP": "*"}, {"POS": "PROPN", "OP": "*"}]
    matcher.add("PROPER_NOUNS", [pattern], greedy="LONGEST")
    doc = nlp(text)

    jobs = []
    matches = matcher(doc)

    for match in matches:
        check_role = str(doc[match[1] : match[2]]).lower()

        for role in job_roles:
            if role in check_role.split() and role != check_role:
                # print(check_role)
                jobs.append(check_role)

    return list(set(jobs))


# print(extract_prev_job_roles(sections))



# Surya dummy

# from summary_text_extractor import extract_prev_job_roles
import re

def tocheckbesideDates(lst):
  nlst=[]
  for i in range(len(lst)-1):
    if lst[i+1][0] - lst[i][-1]<10:
      nlst.append([lst[i],lst[i+1]])
  if len(nlst)>0:
    return nlst[0]
  else:
    return [None,None]
def tocheckNearRole(dateslst,roleslst,text):
  text = text.lower()
  nearest_word = None
  min_distance = float('inf')  # Initialize with infinity
  for word in roleslst:
      distance = abs(text.find(word) - dateslst[0][0])
      if distance < min_distance:
          min_distance = distance
          nearest_word = word
#   print(nearest_word)
  return nearest_word, min_distance

def toExtract_Experience_Rolewise(text):
  i=[]
  j=[]
  # or doc.ents[entity].label_ == "ORG"
  value,dates = ExtractDates(text)
  roles = extract_prev_job_roles(text)
  # toExtract_location(text)
  if len(roles) ==0:
    # print(text)
    pass
  if len(value)>0:
    # print(value)
    for i in range(len(value)-1):
      dat1 = value[i]
      dat2 = value[i+1]
      lst=[]
      for k,j in dates.items():
        if j==dat1 or j==dat2:
          lst.append(k)
      string1 , string2 = tocheckbesideDates(lst)
      if (string1!=None and string2!=None ) and len(roles)>0:
        role,distance = tocheckNearRole([string1,string2],roles,text)
        # print(dat1, "--", dat2)
        # print(role)
  else:
    # print(text)
    pass
from dateutil.parser import parse
from datetime import datetime
def ExtractDates(text):
  # NOt USing date_pattern = r"(?:\d{1,2}[ -]?)?(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\b[ -]\d{1,2}[ -]?\d{2,4}"
  # Using But date_pattern = r"\b(Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember))\s+(\d{2,4})\b\s*(.*?)\s*(?:\b(Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember))\s+(\d{2,4})\b|(\bPresent|Till|Current\b))\b"
  date_pattern =  r"\b(?:\b(Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sept(?:ember)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)\b|\d{2})(?:\b[\s\,/-]+?)(\d{2,4})\b|(\bPresent|Till|Current(?:ly)?|Now\b))\b"
  dates_iter =   re.finditer(date_pattern, text , re.IGNORECASE)
  potential_dates = {datevalue.span():"".join(datevalue.group()) for datevalue in dates_iter}
  # potential_dates = [" ".join(datevalue) for datevalue in re.findall(date_pattern, text , re.IGNORECASE)]
  # print(potential_dates)
  # print(text)
  parsed_dates = []
  for key, date_str in potential_dates.items():
      try:
        if date_str.lower().strip() in ["present","till","current","currently"]:
          date_str = datetime.now().isoformat()
        elif len(date_str.strip().split(" ")[-1])==2:
             date_str = date_str.strip().split(" ")
             date_str[-1] = "20"+date_str[-1]
             date_str = " ".join(date_str)
        parsed_date = parse(date_str, fuzzy=True)
        potential_dates[key]=parsed_date.strftime("%Y-%m-%d")
        parsed_dates.append(parsed_date.strftime("%Y-%m-%d"))
      except ValueError:
          # Skip if parsing fails
          pass
  # Print the parsed dates
  return sorted(list(set(parsed_dates)),reverse= True),potential_dates
  # for date in parsed_dates:
  #     print(date.strftime("%Y-%m-%d %H:%M:%S"))


from tika import parser
def tika_text_extraction(file_path):   
    parsed_file = parser.from_file(file_path)
    text = parsed_file["content"]
    return text


# filename = 'resume.pdf'

# filepath = f"../data/resumes/{filename}";
# text = tika_text_extraction(filepath)

# val = toExtract_Experience_Rolewise(text)




#a dummy end




# try:

#     filename = 'resume2.pdf'

#     filepath = f"../data/resumes/{filename}";

#     # textpdf = extract_text_from_pdf(filepath)
#     text =tika_text_extraction(filepath)
#     print('raw text: ',text)

#     sections, first = divide_resume_sections(text)
#     name = 'Yash'

#     # name = process_get_name(filepath,textpdf,text)

#     newSum = extract_summary(first.lower(), name.lower())
    
#     first60_newSum = extract_first_60_words(newSum)
#     print('\n\n Newly Extracted Summary\n',first60_newSum, '\n New Summary End \n\n' )

    

#     text = tika_text_extraction(f"../data/resumes/{filename}")
    
#     # first_added_element = list(sections.items())[0]
#     print('\n first item \n', first, '\n First End \n')

#     first_paragraph = extract_first_60_words(sections["Summary"])

#     val = toExtract_Experience_Rolewise(sections['Experience'])

#     print('\n\nMoon Dummy\n',val ,'\n\n')



#     print(f'\n Summary for ${filename} \n',first_paragraph,'\n\n')
# except Exception as e:
#     print(f"a exception {e} occured at: {filename}")


# for section_name, content in sections.items():
#     print(f"--- {section_name} ---")
#     print(content)
#     print()


def education_extractor(text):
    desc_col = text
    # desc_col = tika_text_extraction(text)
    
    bachelors = [re.findall("(?<![A-Z])B\.?S\.?c?(?![A-Z])|(?<![A-Z])B\.?A\.?(?![A-Z])|BACHELOR|UNDERGRAD.{0,40} DEGREE|ASSOCIATE'?S?.{20}DEGREE",i, re.IGNORECASE) for i in desc_col.values]
    mba = [re.findall("([\s|-|/]MBA[\s|-|/]|[\s|-|/]MBUS[\s|-|/]|[\s|-|/]MBS[\s|-|/]|MASTERS? OF BUSINESS)",i,re.IGNORECASE) for i in desc_col.values]
    masters = [re.findall("(MASTER'?S?.{0,40}DEGREE|GRADUATE.{0,40}DEGREE|(?<![A-Z])M\.?S\.?(?![A-Z]|\sDYNAMICS|,\sDSC)(?!-?~?\s?OFFICE|\sEXCEL|\sWORD|\sACCESS|-?\s?SQL)|ADVANCED?.{0,40}DEGREE)",i,re.IGNORECASE) for i in desc_col.values]
    phd = [re.findall("(PH\.?D|ADVANCED?.{0,40}DEGREE|DOCTORA[TE|L]|POST-?\s?GRADUATE)",i,re.IGNORECASE) for i in desc_col.values]


import re

# Your regular expressions with raw string literals
bachelors_regex = r"(?<![A-Z])B\.?S\.?c?(?![A-Z])|(?<![A-Z])B\.?A\.?(?![A-Z])|BACHELOR|UNDERGRAD.{0,40} DEGREE|ASSOCIATE'?S?.{20}DEGREE"
mba_regex = r"([\s|-|/]MBA[\s|-|/]|[\s|-|/]MBUS[\s|-|/]|[\s|-|/]MBS[\s|-|/]|MASTERS? OF BUSINESS)"
masters_regex = r"(MASTER'?S?.{0,40}DEGREE|GRADUATE.{0,40}DEGREE|(?<![A-Z])M\.?S\.?(?![A-Z]|\sDYNAMICS|,\sDSC)(?!-?~?\s?OFFICE|\sEXCEL|\sWORD|\sACCESS|-?\s?SQL)|ADVANCED?.{0,40}DEGREE)"
phd_regex = r"(PH\.?D|ADVANCED?.{0,40}DEGREE|DOCTORA[TE|L]|POST-?\s?GRADUATE)"

# Example text input
text_input = "I have a BSc in Computer Science and an MBA. I also completed my Master's degree and PhD."

# Function to match text against the provided regular expressions
def match_education(text):
    bachelors_matches = re.findall(bachelors_regex, text, re.IGNORECASE)
    mba_matches = re.findall(mba_regex, text, re.IGNORECASE)
    masters_matches = re.findall(masters_regex, text, re.IGNORECASE)
    phd_matches = re.findall(phd_regex, text, re.IGNORECASE)
    
    return {
        "Bachelors": bachelors_matches,
        "MBA": mba_matches,
        "Masters": masters_matches,
        "PhD": phd_matches
    }

# Match the text against the regular expressions
matches = match_education('../data/resumes/resume2.pdf')

# Display the matches
for degree, matches_list in matches.items():
    if matches_list:
        print(f"{degree}: {matches_list}")
