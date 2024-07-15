from pdfminer.converter import TextConverter
from pdfminer.pdfinterp import PDFPageInterpreter
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from docx import Document

import spacy
from spacy.matcher import Matcher
import io
import requests
import re

def extract_text_from_pdf(pdf_path):

    
    if pdf_path.startswith('http'):
        response = requests.get(pdf_path)
        pdf_data = response.content
    else:
        with open(pdf_path, 'rb') as fh:
            pdf_data = fh.read()

    with io.BytesIO(pdf_data) as fh:
        for page in PDFPage.get_pages(fh, caching=True, check_extractable=True):
            resource_manager = PDFResourceManager()
            fake_file_handle = io.StringIO()
            converter = TextConverter(resource_manager, fake_file_handle, laparams=LAParams())
            page_interpreter = PDFPageInterpreter(resource_manager, converter)
            page_interpreter.process_page(page)
            text = fake_file_handle.getvalue()
            yield text
            converter.close()
            fake_file_handle.close()

def extract_from_docx(docx_file_path):
    doc = Document(docx_file_path)
    text = []
    for para in doc.paragraphs:
        text.append(para.text)

    tables_data = []
    for table in doc.tables:
        table_data = []  # List to hold all rows in a table
        for row in table.rows:
            row_data = []  # List to hold cell values in a row
            for cell in row.cells:
                row_data.append(cell.text)
            table_data.append(row_data)  # Append row data to table data
        tables_data.append(table_data)  # Append table data to tables data

        # Add table data to text
        table_text = '\n'.join(['\t'.join(row) for row in table_data])
        text.append(table_text)

    # return '\n'.join(text), tables_data
    return "\n".join(text)


from tika import parser
def tika_text_extraction(file_path):   
    parsed_file = parser.from_file(file_path)
    text = parsed_file["content"]
    return text


def divide_resume_sections(resume_text):
    sections = {}
  
    section_patterns = {

        "Summary":r"ABOUT|Overview|OVERVIEW|Summary|SUMMARY|summary|overview",
        "Education": r"Education|EDUCATION",
        "Experience": r"\b(Experience|EXPERIENCE|Training|TRAINING|TRAININGS|INTERNSHIPS)\b",
        # "Training":r"Training|TRAINING|TRAININGS|INTERNSHIPS|INTERNSHIPS"        
        "Projects": r"Projects|PROJECTS",
        "Skills":r"SKILLS",
        "Achievements":"Achievements|ACHIEVEMENTS"
     
    }
    for sec in section_patterns:
        sections[sec]=""


    sect = {"Summary":["ABOUT","Overview","OVERVIEW","Summary","SUMMARY","summary","overview"],"Education":["Education","EDUCATION"],"Experience":["Experience","EXPERIENCE","Training","TRAINING","TRAININGS","INTERNSHIPS"],"Projects":["Projects","PROJECTS"],"Skills":["SKILLS",""],"Achievements":["Achievements","ACHIEVEMENTS"]}
    
    for section_name in section_patterns:       
        matches = re.search(section_patterns[section_name], resume_text)       
        if matches:
            # print(matches.group(0))
            section_start = matches.start()      
            next_section_start = re.search(r"\b(ABOUT|Overview|OVERVIEW|Summary|SUMMARY|Education|EDUCATION|Experience|EXPERIENCE|Training|TRAINING|TRAININGS|INTERNSHIPS|Projects|PROJECT|SKILLS|Achievements|ACHIEVEMENTS|$)\b", resume_text[section_start+1:])          
            count=0
            check_end2 = 0
            final = 0

            # print(sec[section_name])
            if next_section_start:
                # print("xhexk ", next_section_start, next_section_start.group(0))
                while sect[section_name][0]==next_section_start.group(0) or sect[section_name][1]==next_section_start.group(0) :
                    check_end2 += section_start + next_section_start.end() +1 
                    next_section_start = re.search(r"\b(ABOUT|Overview|OVERVIEW|Summary|SUMMARY|Education|EDUCATION|Experience|EXPERIENCE|Training|TRAINING|TRAININGS|INTERNSHIPS|Projects|PROJECTS|SKILLS|Achievements|ACHIEVEMENTS|$)\b", resume_text[check_end2:])    
                    if next_section_start:           
                        if sect[section_name][0]!=next_section_start.group(0) or sect[section_name][1]!=next_section_start.group(0):
                            final = next_section_start.start()+check_end2    
                    else:
                        break            
                    count+=1
                    if count==5:
                        break
                    flag=1  
            
            if next_section_start and next_section_start.group(0).replace(" ",""):                          
                section_end = section_start + next_section_start.start() +1              
            else:                
                section_end = len(resume_text)        
            if section_name in sections:                
                sections[section_name] += resume_text[section_start:section_end].strip()
            else:                
                sections[section_name] = resume_text[section_start:section_end].strip()

    return sections


# text=""
# for page in extract_text_from_pdf("data/resumes/resume.pdf"):
#     text += ' ' + page 
# print(text)
# print(tika_text_extraction("data/resumes/resume.pdf"))

# text ="""A data-driven Product Manager who has a proven track record of boosting revenue with 9 years of experience.                                        
# Experienced in tools like LLM, Midjourney, Stable diffusion, JIRA, Mixpanel, Google Analytics, A/B or multi 
# variate testing. On the business front corporate finance and Pricing strategy are my forte. Led the cross-
# functional teams in B2B, SaaS, B2C, marketing, sales, and service. Developed and demonstrated go-to-market 
# strategies, digital strategies, and product roadmaps. Successful KPIs include raising engagement levels, 
# enhancing client satisfaction, and generating revenue through efficient tactics."""
# print(text)
# sections = divide_resume_sections(text)["Summary"]
# print(sections)
# # for section_name, content in sections.items():
# #     print(f"--- {section_name} ---")
# #     print(content)
# #     print()

# nlp = spacy.load("en_core_web_sm")

# matcher = Matcher(nlp.vocab)

# def extract_prev_job_roles(text):
#     job_roles = [
#         "intern",
#         "internship" "analyst",
#         "developer",
#         "manager",
#         "engineer",
#         "consultant",
#         "designer",
#         "specialist",
#         "coordinator",
#         "administrator",
#         "executive",
#         "assistant",
#         "supervisor",
#         "technician",
#         "associate",
#         "officer",
#         "leader",
#         "expert",
#         "advisor",
#         "strategist",
#         "resources",
#         "test",
#     ]

#     text = text.replace("\n", " ")
#     # print(text)
#     matcher = Matcher(nlp.vocab)
#     pattern = [{"POS": "NOUN", "OP": "*"}, {"POS": "PROPN", "OP": "*"}]
#     matcher.add("PROPER_NOUNS", [pattern], greedy="LONGEST")
#     doc = nlp(text)

#     jobs = []
#     matches = matcher(doc)

#     for match in matches:
#         check_role = str(doc[match[1] : match[2]]).lower()

#         for role in job_roles:
#             if role in check_role.split() and role != check_role:
#                 # print(check_role)
#                 jobs.append(check_role)

#     return list(set(jobs))


# print(extract_prev_job_roles(sections))